import docx
import os
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Cm, Inches


def to_five(num):
    if len(num) == 5:
        return num
    else:
        a = "0" * (5 - len(num))
        return a + num


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None



first_num = input("Введите номер первого талона:\n")
fn = first_num
eat_time = input("Какой прием пищи? (введите цифру):\nЗавтрак - 1, Обед - 2, Ужин - 3\n")
if eat_time == "1":
    eat_time = "Завтрак"
elif eat_time == "2":
    eat_time = "Обед"
else:
    eat_time = "Ужин"

doc = docx.Document("layout.docx")
#tab = doc.add_table(rows=6,cols=4, style="Table Grid")

tab = doc.tables[0]
for row in range(48):
    for col in range(4):
        cell = tab.cell(row, col)
        cell.width = Inches(5)



        #p1 = cell.add_paragraph('ИП Сивобород М.С.')
        #tab.cell(row,col).paragraphs[0].paragraph_format
        run = cell.add_paragraph().add_run('ИП Иванов И.И.')
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(12)
        
        run = cell.add_paragraph().add_run('«БАЙКАЛ-ВЕСТ»')
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(18)
        
        run = cell.add_paragraph().add_run('Талон на питание')
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(16)

        run = cell.add_paragraph().add_run(to_five(first_num))
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(18)

        first_num = str(int(first_num) + 1)

        run = cell.add_paragraph().add_run(eat_time)
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(18)
        
        run = cell.add_paragraph().add_run('Дата__________')
        font = run.font 
        font.name = 'Calibri'
        font.size = Pt(18)

        g = cell.paragraphs[0]
        delete_paragraph(g)
        for i in range(6):
            tab.cell(row,col).paragraphs[i].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

homeDir = os.path.expanduser('~')
print(homeDir + r'\Desktop')       # Example: C:\Users\Username\Desktop
path = homeDir + r'\Desktop' + '\\' + fn + '-' + first_num + '_' + eat_time + '.docx'
doc.save(path)